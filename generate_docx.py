import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Add Title
title = doc.add_heading('Payment Flow (Customer Journey) - LK Tech', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add Introduction
doc.add_paragraph(
    'Dokumen ini menjelaskan alur end-to-end proses pemesanan barang/jasa sampai '
    'dengan proses pembayaran pada website LK Tech (lingkungan pengembangan/localhost), '
    'sebagai persyaratan verifikasi akun Midtrans.'
)

# ---- DIAGRAM SECTION ----
doc.add_heading('User Flow Diagram', level=1)
diagram_intro = doc.add_paragraph('Berikut adalah diagram alur (customer journey) dari awal hingga selesai:')

# Create a visual flow using nested lists or styled text
flow_steps = [
    ("1. Halaman Beranda", "Customer mengunjungi website LK Tech"),
    ("2. Halaman Produk & Keranjang", "Memilih produk dan menambahkannya ke keranjang"),
    ("3. Halaman Checkout", "Mengisi data diri dan detail alamat"),
    ("4. Pembayaran (Midtrans Snap)", "Klik bayar dan pop-up Midtrans muncul"),
    ("5. Proses Transfer/Bayar", "Pilih metode (VA/GoPay/dll) dan lakukan pembayaran"),
    ("6. Transaksi Berhasil", "Diarahkan ke halaman sukses dan menerima Invoice")
]

for i, (step_title, step_desc) in enumerate(flow_steps):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a box-like effect using borders? Python-docx doesn't easily do paragraph borders, 
    # but we can use highlight/bold text to make it stand out.
    run = p.add_run(f"[ {step_title} ]")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 112, 192) # Blue color
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(step_desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100) # Gray color
    
    if i < len(flow_steps) - 1:
        # Add arrow down
        p_arrow = doc.add_paragraph()
        p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_arrow = p_arrow.add_run("⬇")
        run_arrow.font.size = Pt(20)
        run_arrow.bold = True

doc.add_page_break()

# ---- DETAIL SECTION ----
doc.add_heading('Detail Langkah & Screenshot', level=1)

# Step 1
p = doc.add_paragraph()
p.add_run('1. Halaman Beranda (Home Page)').bold = True
doc.add_paragraph('Customer membuka situs web LK Tech dan melihat daftar produk atau layanan IT yang ditawarkan.', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT HALAMAN HOME DI SINI]', style='Intense Quote')

# Step 2
p = doc.add_paragraph()
p.add_run('2. Halaman Produk & Keranjang Belanja (Add to Cart)').bold = True
doc.add_paragraph('Customer memilih produk/layanan yang diinginkan dan menambahkannya ke keranjang belanja.', style='List Bullet')
doc.add_paragraph('Customer meninjau pesanan di halaman keranjang belanja sebelum melanjutkan.', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT HALAMAN KERANJANG/PRODUK DI SINI]', style='Intense Quote')

# Step 3
p = doc.add_paragraph()
p.add_run('3. Halaman Checkout (Pengisian Data)').bold = True
doc.add_paragraph('Customer mengisi form detail penagihan (Billing Details) seperti nama, email, nomor telepon, dan alamat.', style='List Bullet')
doc.add_paragraph('Sistem menampilkan ringkasan pesanan beserta total biaya yang harus dibayar.', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT HALAMAN CHECKOUT/FORM DATA DIRI DI SINI]', style='Intense Quote')

# Step 4
p = doc.add_paragraph()
p.add_run('4. Proses Pembayaran (Payment Gateway - Midtrans Snap)').bold = True
doc.add_paragraph('Customer mengklik tombol "Bayar" atau "Place Order".', style='List Bullet')
doc.add_paragraph('Sistem memunculkan halaman pembayaran (Payment Page / Snap Pop-up) dari Midtrans.', style='List Bullet')
doc.add_paragraph('Halaman ini menampilkan metode pembayaran yang tersedia (misal: Virtual Account BCA, Mandiri, e-Wallet GoPay, QRIS, dll).', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT HALAMAN METODE PEMBAYARAN MIDTRANS DI SINI]', style='Intense Quote')

# Step 5
p = doc.add_paragraph()
p.add_run('5. Instruksi & Penyelesaian Pembayaran').bold = True
doc.add_paragraph('Customer memilih salah satu metode pembayaran dan sistem menampilkan instruksi pembayaran (contoh: nomor Virtual Account dan cara transfer).', style='List Bullet')
doc.add_paragraph('Customer melakukan pembayaran sesuai dengan instruksi tersebut.', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT INSTRUKSI PEMBAYARAN MIDTRANS DI SINI]', style='Intense Quote')

# Step 6
p = doc.add_paragraph()
p.add_run('6. Halaman Sukses & Invoice (Payment Success)').bold = True
doc.add_paragraph('Setelah pembayaran berhasil diselesaikan, Midtrans memberikan notifikasi otomatis ke sistem LK Tech.', style='List Bullet')
doc.add_paragraph('Customer diarahkan ke halaman "Payment Success" atau halaman Terima Kasih.', style='List Bullet')
doc.add_paragraph('Sistem mengirimkan invoice atau tanda terima ke email customer.', style='List Bullet')
doc.add_paragraph('[SISIPKAN CAPTURE/SCREENSHOT HALAMAN SUCCESS / INVOICE DI SINI]', style='Intense Quote')

# Save document
output_path = 'Payment_Flow_LK_Tech.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
