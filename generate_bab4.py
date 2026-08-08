"""
Script Generate BAB 4 - Mursanto (19252417)
BAB IV: PENUTUP
4.1 Kesimpulan
4.2 Saran
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

DST = r'D:\Project\Kebutuhan Skripsi\SKRIPSI\BAB 4_MURSANTO_19252417.docx'

# Buat dokumen baru
doc = Document()

# ============================================================
# PAGE SETUP
# ============================================================
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(4.0)
section.right_margin = Cm(3.0)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color='2F4F8F'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_para(text='', bold=False, italic=False, alignment='justify', font_size=12,
             space_before=0, space_after=200, first_line_indent=True):
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    if alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if first_line_indent:
        fmt.first_line_indent = Cm(1.25)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
    return para

def add_heading(text, level=1, font_size=14):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = para.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Cm(0)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    return para

def add_bullet(text, font_size=12, indent_level=1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.left_indent = Cm(indent_level * 1.25)
    fmt.first_line_indent = Cm(-0.5)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(100)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    return para

def add_numbered_item(number, text, font_size=12, indent=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    if indent:
        fmt.left_indent = Cm(1.25)
        fmt.first_line_indent = Cm(-1.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(120)
    run_num = para.add_run(f'{number}. ')
    run_num.bold = True
    run_num.font.size = Pt(font_size)
    run_num.font.name = 'Times New Roman'
    run_text = para.add_run(text)
    run_text.font.size = Pt(font_size)
    run_text.font.name = 'Times New Roman'
    return para

# ============================================================
# BAB IV HEADER
# ============================================================

# BAB IV
bab_heading = add_heading('BAB IV', level=1, font_size=14)

# PENUTUP
penutup_heading = add_heading('PENUTUP', level=1, font_size=14)

doc.add_paragraph()

# ============================================================
# 4.1 KESIMPULAN
# ============================================================

add_heading('4.1 Kesimpulan', level=2, font_size=12)

intro = add_para(
    'Berdasarkan hasil penelitian, perancangan, implementasi, dan pengujian yang telah dilaksanakan '
    'terhadap Sistem Informasi Inventori dan Penjualan Berbasis Web dengan Peningkatan Keamanan '
    'Autentikasi pada LKtech, dapat ditarik beberapa kesimpulan sebagai berikut:',
    alignment='justify', font_size=12, first_line_indent=True
)
intro.paragraph_format.space_after = Pt(200)

# Kesimpulan 1 - Sistem berhasil dibangun
add_numbered_item(
    1,
    'Sistem informasi inventori dan penjualan berbasis web untuk LKtech telah berhasil dikembangkan '
    'menggunakan framework Laravel dengan PHP 8.3 dan basis data MySQL. Sistem ini berhasil '
    'mengintegrasikan empat modul operasional utama—inventori, penjualan (Point of Sales), penyewaan, '
    'dan servis—ke dalam satu platform terpadu yang dapat diakses secara online melalui domain '
    'lktech.online. Kehadiran sistem ini secara langsung menggantikan proses pencatatan manual '
    'berbasis Microsoft Excel yang sebelumnya digunakan LKtech, sehingga seluruh data operasional '
    'kini terpusat, terintegrasi, dan dapat diakses secara real-time.',
    font_size=12
)

# Kesimpulan 2 - RBAC
add_numbered_item(
    2,
    'Implementasi Role-Based Access Control (RBAC) menggunakan library spatie/laravel-permission '
    'berhasil mewujudkan pemisahan wewenang operasional yang terstruktur di antara tiga hierarki '
    'pengguna: Super Admin, Kasir, dan Teknisi. Setiap peran memiliki hak akses yang dikontrol '
    'secara ketat oleh middleware Laravel, sehingga setiap pengguna hanya dapat mengakses modul '
    'dan fungsionalitas yang sesuai dengan tanggung jawab kerjanya. Upaya akses ke rute yang '
    'tidak diizinkan secara otomatis menghasilkan respons penolakan 403, memastikan tidak ada '
    'penyalahgunaan wewenang dalam lingkungan sistem.',
    font_size=12
)

# Kesimpulan 3 - 2FA
add_numbered_item(
    3,
    'Peningkatan keamanan autentikasi melalui implementasi Two-Factor Authentication (2FA) '
    'berbasis Time-Based One-Time Password (TOTP) menggunakan Google Authenticator berhasil '
    'diterapkan secara wajib (mandatory) bagi seluruh pengguna internal sistem. Middleware '
    'Ensure2faSetup.php memastikan bahwa pengguna dengan peran Admin, Kasir, dan Teknisi '
    'diwajibkan mengaktifkan 2FA sebelum dapat mengakses sistem. Mekanisme ini secara signifikan '
    'meningkatkan perlindungan sistem dari ancaman pencurian kredensial (credential theft), '
    'serangan brute-force, dan akses tidak sah, karena penyerang memerlukan kode OTP 6-digit yang '
    'berubah setiap 30 detik di samping kata sandi yang benar.',
    font_size=12
)

# Kesimpulan 4 - Keamanan berlapis
add_numbered_item(
    4,
    'Sistem mengimplementasikan keamanan berlapis yang komprehensif mencakup: hash password Bcrypt '
    'melalui fungsi Hash::make() untuk melindungi kredensial dari kebocoran basis data; '
    'proteksi CSRF melalui middleware VerifyCsrfToken dan direktif @csrf pada setiap form; '
    'proteksi XSS melalui Blade Template Engine yang melakukan HTML Entity Encoding otomatis; '
    'serta pencegahan SQL Injection melalui Eloquent ORM dengan PDO Prepared Statements. '
    'Kombinasi lapisan keamanan ini menjadikan sistem LKtech memenuhi standar keamanan aplikasi '
    'web modern yang diperlukan untuk melindungi data operasional UMKM.',
    font_size=12
)

# Kesimpulan 5 - Activity Log
add_numbered_item(
    5,
    'Fitur Activity Log berhasil diimplementasikan sebagai instrumen audit keamanan yang merekam '
    'secara otomatis seluruh aktivitas pengguna dalam sistem, termasuk waktu login, jenis aksi '
    'yang dilakukan (tambah, ubah, hapus data), informasi detail tentang data yang dimodifikasi, '
    'serta alamat IP dan informasi perangkat yang digunakan. Data Activity Log yang hanya dapat '
    'diakses oleh Super Admin ini menjadi fondasi investigasi insiden keamanan dan pemantauan '
    'perilaku pengguna secara transparan.',
    font_size=12
)

# Kesimpulan 6 - Pengujian
add_numbered_item(
    6,
    'Berdasarkan hasil pengujian Black-Box Testing yang dilaksanakan terhadap 13 skenario kritis, '
    'seluruh fungsionalitas sistem—mulai dari proses autentikasi 2FA, pengelolaan inventori, '
    'transaksi Point of Sales dengan auto-deduct stok, manajemen penyewaan, pencatatan tiket '
    'servis, hingga validasi hak akses berbasis peran—memberikan keluaran yang sesuai dengan '
    'hasil yang diharapkan. Tidak ditemukan satu pun skenario yang menghasilkan keluaran '
    'menyimpang dari spesifikasi, sehingga sistem dinyatakan lolos uji dan siap untuk '
    'dioperasikan di lingkungan produksi.',
    font_size=12
)

doc.add_paragraph()

# ============================================================
# TABEL PENCAPAIAN TUJUAN
# ============================================================

caption_tbl = add_para(
    'Tabel 4.1 Pencapaian Tujuan Penelitian',
    bold=True, alignment='center', font_size=11, first_line_indent=False
)
caption_tbl.paragraph_format.space_after = Pt(100)

pencapaian_data = [
    ('1',
     'Mengembangkan sistem informasi inventori dan penjualan berbasis web untuk LKtech menggunakan framework Laravel.',
     'Sistem berhasil dibangun dengan Laravel + PHP 8.3 + MySQL, live di lktech.online.',
     'Tercapai'),
    ('2',
     'Mengembangkan halaman publik (landing page) terintegrasi sebagai platform penjualan langsung (direct selling).',
     'Landing page dinamis berhasil diimplementasikan, menampilkan katalog produk real-time tanpa perlu login.',
     'Tercapai'),
    ('3',
     'Menerapkan Role-Based Access Control (RBAC) untuk memisahkan wewenang antara Super Admin, Kasir, dan Teknisi.',
     'RBAC dengan spatie/laravel-permission berhasil diterapkan; uji akses tidak sah menghasilkan respon 403.',
     'Tercapai'),
    ('4',
     'Mengimplementasikan keamanan Two-Factor Authentication (2FA) menggunakan Google Authenticator.',
     '2FA mandatory berhasil diimplementasikan; seluruh pengguna internal wajib mengaktifkan 2FA sebelum akses sistem.',
     'Tercapai'),
]

headers_tbl = ['No', 'Tujuan Penelitian', 'Hasil Implementasi', 'Status']
col_widths_tbl = [Cm(0.8), Cm(5.5), Cm(7.5), Cm(2.0)]

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Normal Table'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate(col_widths_tbl):
    tbl.columns[i].width = w

hdr = tbl.rows[0].cells
for i, (cell, header) in enumerate(zip(hdr, headers_tbl)):
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(cell, '1a3c5e')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Times New Roman'

for no, tujuan, hasil, status in pencapaian_data:
    row = tbl.add_row().cells
    vals = [no, tujuan, hasil, status]
    for j, (cell, val) in enumerate(zip(row, vals)):
        cell.text = val
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
        if j == 3:
            set_cell_background(cell, 'D5F5E3')
        elif int(no) % 2 == 0:
            set_cell_background(cell, 'EEF2FF')

set_table_borders(tbl, 'B0BEC5')

src_note = add_para(
    '(Sumber: Olahan Peneliti, 2026)',
    italic=True, alignment='center', font_size=10, first_line_indent=False
)

doc.add_paragraph()

# ============================================================
# 4.2 SARAN
# ============================================================

add_heading('4.2 Saran', level=2, font_size=12)

saran_intro = add_para(
    'Meskipun sistem informasi LKtech telah berhasil dibangun dan dinyatakan lolos pengujian, '
    'terdapat beberapa rekomendasi perbaikan dan pengembangan lanjutan yang perlu dipertimbangkan '
    'untuk meningkatkan kapabilitas, keamanan, dan keberlanjutan sistem di masa mendatang:',
    alignment='justify', font_size=12, first_line_indent=True
)
saran_intro.paragraph_format.space_after = Pt(200)

# Saran 1 - Auto Logout
add_numbered_item(
    1,
    'Implementasi Mekanisme Auto-Logout Berbasis Idle Session. '
    'Saat ini sistem belum memiliki fitur yang memaksa pengguna keluar (logout) secara otomatis '
    'ketika sesi dibiarkan tidak aktif (idle) dalam jangka waktu tertentu. Kondisi ini menimbulkan '
    'risiko keamanan, khususnya pada PC Kasir yang sering ditinggalkan di area publik. Disarankan '
    'agar pengembang selanjutnya menambahkan konfigurasi session timeout pada file config/session.php '
    'dengan durasi maksimal 10–15 menit idle, disertai notifikasi peringatan kepada pengguna '
    'sebelum sesi diakhiri.',
    font_size=12
)

# Saran 2 - Dashboard Monitoring Log
add_numbered_item(
    2,
    'Pengembangan Dashboard Monitoring Activity Log yang Lebih Interaktif. '
    'Fitur Activity Log yang telah diimplementasikan saat ini menampilkan data dalam format tabel '
    'sederhana. Disarankan untuk mengembangkan antarmuka monitoring yang lebih interaktif dengan '
    'menambahkan visualisasi grafis (chart) tren aktivitas, filter pencarian berdasarkan pengguna '
    'dan jenis aksi, serta sistem notifikasi otomatis (email/push notification) kepada Super Admin '
    'ketika terdeteksi anomali mencurigakan seperti percobaan login berulang yang gagal '
    '(indikasi serangan brute-force).',
    font_size=12
)

# Saran 3 - Backup Otomatis
add_numbered_item(
    3,
    'Implementasi Sistem Backup Database Terjadwal (Automated Backup). '
    'Sistem saat ini bergantung pada proses backup manual yang dilakukan secara insidental. '
    'Untuk meningkatkan ketahanan sistem terhadap kegagalan server (disaster recovery), '
    'disarankan untuk mengimplementasikan CRON job terjadwal yang secara otomatis mencadangkan '
    'dump database MySQL ke layanan cloud storage (misalnya Google Drive atau Amazon S3) '
    'minimal satu kali sehari pada jam operasional rendah. Hal ini memastikan pemulihan data '
    'dapat dilakukan dengan cepat jika terjadi kegagalan infrastruktur.',
    font_size=12
)

# Saran 4 - SOP dan Dokumentasi
add_numbered_item(
    4,
    'Penyusunan Standard Operating Procedure (SOP) Keamanan Sistem. '
    'Meskipun sistem telah dilengkapi dengan mekanisme keamanan teknis yang memadai, '
    'investasi teknologi ini perlu diimbangi dengan pembentukan SOP administratif yang '
    'mendokumentasikan kewajiban dan tanggung jawab setiap pengguna terkait pengelolaan '
    'kredensial akses, prosedur aktivasi 2FA bagi karyawan baru, kebijakan penggantian '
    'password secara berkala, serta prosedur tanggap darurat ketika terjadi insiden '
    'keamanan. SOP ini menjadi landasan budaya keamanan informasi di tingkat organisasi.',
    font_size=12
)

# Saran 5 - Fitur Notifikasi Stok
add_numbered_item(
    5,
    'Penambahan Fitur Notifikasi Stok Minimum (Low Stock Alert). '
    'Sistem saat ini menampilkan data stok secara real-time namun belum memiliki mekanisme '
    'peringatan otomatis ketika stok suatu produk mendekati batas minimum. Disarankan untuk '
    'menambahkan fitur low stock alert yang mengirimkan notifikasi kepada Super Admin dan Kasir '
    'melalui antarmuka sistem maupun email ketika kuantitas produk tertentu berada di bawah '
    'ambang batas yang telah ditentukan, sehingga proses pengadaan barang dapat dilakukan '
    'secara proaktif sebelum stok habis.',
    font_size=12
)

# Saran 6 - Laporan Keuangan Lebih Lengkap
add_numbered_item(
    6,
    'Pengembangan Modul Laporan Keuangan yang Lebih Komprehensif. '
    'Modul laporan yang ada saat ini mencakup laporan penjualan harian dan bulanan. '
    'Untuk meningkatkan nilai strategis sistem bagi manajemen LKtech, disarankan '
    'untuk mengembangkan modul laporan keuangan yang lebih komprehensif mencakup '
    'laporan laba rugi (profit & loss statement) yang terperinci, analisis tren penjualan '
    'per kategori produk, laporan utilisasi unit sewa, serta proyeksi pendapatan berbasis '
    'data historis. Laporan-laporan ini dapat disajikan dalam format yang dapat diekspor '
    'ke PDF dan Excel untuk keperluan manajemen.',
    font_size=12
)

# Saran 7 - Pengembangan Aplikasi Mobile
add_numbered_item(
    7,
    'Pertimbangan Pengembangan Aplikasi Mobile (Progressive Web App). '
    'Sistem saat ini bersifat web-based dan dapat diakses melalui browser di perangkat apapun. '
    'Untuk meningkatkan fleksibilitas akses bagi Teknisi yang sering berpindah lokasi saat '
    'menangani servis di luar toko, disarankan untuk mengembangkan Progressive Web App (PWA) '
    'atau aplikasi mobile ringan yang memungkinkan akses dan pembaruan status tiket servis '
    'secara offline, kemudian melakukan sinkronisasi data ketika koneksi internet tersedia kembali.',
    font_size=12
)

doc.add_paragraph()

# ============================================================
# PENUTUP PARAGRAF
# ============================================================
closing = add_para(
    'Demikian kesimpulan dan saran yang dapat penulis sampaikan berdasarkan hasil penelitian '
    'ini. Penulis berharap bahwa sistem informasi yang telah dikembangkan dapat memberikan '
    'manfaat yang nyata bagi operasional LKtech dan menjadi kontribusi yang berarti bagi '
    'pengembangan keilmuan di bidang sistem informasi, khususnya dalam hal integrasi '
    'mekanisme keamanan autentikasi pada sistem informasi bisnis berbasis web skala UMKM. '
    'Penelitian ini diharapkan pula dapat menjadi referensi bagi peneliti-peneliti berikutnya '
    'yang ingin mengembangkan sistem serupa dengan cakupan fungsionalitas yang lebih luas '
    'dan mekanisme keamanan yang lebih komprehensif.',
    alignment='justify', font_size=12, first_line_indent=True
)

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f"SUKSES! BAB 4 berhasil disimpan ke:\n   {DST}")
print("\nStruktur BAB 4:")
print("  4.1 Kesimpulan")
print("      - 6 poin kesimpulan terstruktur")
print("      - Tabel 4.1 Pencapaian Tujuan Penelitian")
print("  4.2 Saran")
print("      - 7 saran pengembangan lanjutan")
print("      - Paragraf penutup")
