from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = r'D:\Project\Kebutuhan Skripsi\Arsitektur Enterprise\Laporan_Audit_COBIT2019_LKTech.docx'

doc = Document()

# Page setup
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_para(text='', bold=False, italic=False, alignment='left', font_size=12, space_after=120):
    para = doc.add_paragraph()
    if alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
    return para

def add_heading(text, level=1):
    para = doc.add_paragraph()
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_size = 16
    elif level == 2:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font_size = 14
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font_size = 12
    
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    return para

# Title Page
add_heading('LAPORAN HASIL AUDIT SISTEM INFORMASI', level=1)
add_heading('Sistem Informasi Manajemen, Inventori, dan Penjualan (LKTech)', level=1)

add_para()
add_para('Organisasi: LKTech', bold=True)
add_para('Periode Audit: 1 Juni 2026 – 29 Juni 2026', bold=True)
add_para('Versi Dokumen: v1.0', bold=True)

add_para()
add_para('Disusun oleh:', bold=True)
add_para('• 19251913 - Catur Yudha Yunandar\n• 19252417 - Mursanto\n• 19252397 - Hafid Robbi A\n• 19252111 - Imas Wulansari\n• 19251909 - Sarah Nurlianah')

add_para()
add_para('Diverifikasi oleh: Yuri Yuliani, M.Kom', bold=True)
doc.add_page_break()

# DAFTAR ISI
add_heading('DAFTAR ISI', level=1)
daftar_isi = [
    "1. Executive Summary", "2. Latar Belakang & Tujuan", "3. Ruang Lingkup Audit",
    "4. Metodologi Audit", "5. Gambaran Objek Audit (Sisfo)", "6. Hasil Penilaian Capability (Per Proses/Praktik/Aktivitas)",
    "7. Hasil Perhitungan Maturity", "8. Temuan Audit & Analisis Risiko", "9. Analisis Gap (Current vs Target)",
    "10. Rekomendasi & Rencana Tindak Lanjut", "11. Roadmap Peningkatan", "12. Kesimpulan", "13. Lampiran"
]
for item in daftar_isi:
    add_para(item, font_size=12, space_after=6)
doc.add_page_break()

# 1. Executive Summary
add_heading('1. Executive Summary', level=2)
add_para('1.1 Ringkasan Tujuan Audit:', bold=True)
add_para('Menilai kapabilitas proses operasional sistem informasi LKTech berdasarkan kerangka COBIT 2019, secara khusus berfokus pada DSS04 (Managed Continuity) terkait keandalan cadangan data, dan DSS05 (Managed Security Services) terkait kontrol akses serta keamanan autentikasi.', alignment='justify')

add_para('1.2 Ringkasan Ruang Lingkup:', bold=True)
add_para('Proses COBIT yang dinilai mencakup DSS04 dan DSS05. Aktivitas fokus berada pada praktik pencadangan data (backup) serta pengelolaan proteksi endpoint dan sesi pengguna (session timeout, 2FA, RBAC). Periode pengambilan bukti dilakukan sepanjang bulan Juni 2026 pada sistem yang berstatus live di cloud hosting.', alignment='justify')

add_para('1.3 Ringkasan Hasil (Highlight):', bold=True)
add_para('Current capability untuk proses keamanan (DSS05) tergolong baik di Level 3 (Established) berkat implementasi RBAC (Role-Based Access Control) dan 2FA (Two-Factor Authentication). Namun, untuk proses kelangsungan layanan (DSS04), sistem masih berada di Level 2 (Managed) karena pencadangan basis data (backup) masih dilakukan secara manual/insidental. Ditemukan 2 risiko kritis: ketiadaan batasan waktu mati otomatis (auto-logout sesi) dan pencadangan manual.', alignment='justify')

add_para('1.4 Kesimpulan Eksekutif:', bold=True)
add_para('Secara keseluruhan, arsitektur keamanan dasar aplikasi sudah terbangun kokoh. Namun, operasi TI belum sepenuhnya otonom. Rekomendasi utama dan paling mendesak adalah mengaktifkan CRON Job untuk backup otomatis harian ke penyimpanan eksternal, menerapkan session timeout sebesar 10 menit untuk perlindungan PC kasir, serta mengembangkan antarmuka Dashboard Activity Log secara visual untuk mempermudah pemantauan oleh Super Admin.', alignment='justify')

# 2. Latar Belakang & Tujuan
add_heading('2. Latar Belakang & Tujuan', level=2)
add_para('2.1 Latar Belakang:', bold=True)
add_para('LKTech telah mentransformasi layanan operasionalnya dari pencatatan manual menjadi sistem informasi terintegrasi berbasis web (lktech.online). Sistem ini mengelola data sensitif transaksi (Point of Sales) dan inventori bernilai tinggi. Berdasarkan evaluasi arsitektur sebelumnya (menggunakan TOGAF ADM), sistem menunjukkan kerentanan pada tata kelola pengamanan perangkat di area publik (PC Kasir) dan risiko kehilangan data transaksional akibat belum adanya otomasi pemulihan bencana (disaster recovery).', alignment='justify')

add_para('2.2 Tujuan Audit:', bold=True)
add_para('1. Mengukur kapabilitas (Capability Level) dari proses pengelolaan keamanan (DSS05) dan kelangsungan layanan (DSS04) menggunakan kerangka kerja COBIT 2019.\n2. Mengidentifikasi gap atau kesenjangan arsitektural antara kondisi saat ini (as-is) dengan target ideal (to-be).\n3. Merumuskan rencana tindak lanjut taktis (action plan) untuk memperkuat daya tahan sistem (resiliensi) terhadap kegagalan perangkat atau penyalahgunaan akun.', alignment='justify')

add_para('2.3 Kriteria/Standar/Acuan:', bold=True)
add_para('COBIT 2019 Framework (Fokus domain Deliver, Service, and Support / DSS).', alignment='justify')

# 3. Ruang Lingkup Audit
add_heading('3. Ruang Lingkup Audit', level=2)
add_para('3.1 Objek Audit:', bold=True)
add_para('Sistem Informasi Manajemen, Inventori, dan Penjualan (LKTech) berbasis web. Pemilik sistem: Super Admin (Pemilik Usaha). Pengguna utama: Kasir, Teknisi, dan Pelanggan publik. Teknologi utama: Laravel, PHP 8.2, MySQL, disebarkan pada cPanel Cloud Hosting.', alignment='justify')

add_para('3.2 Proses COBIT yang Dinilai:', bold=True)
add_para('• DSS04: Managed Continuity (Manajemen kelangsungan layanan dan backup).\n• DSS05: Managed Security Services (Manajemen proteksi keamanan akses dan sesi pengguna).', alignment='justify')

# 4. Metodologi Audit
add_heading('4. Metodologi Audit', level=2)
add_para('4.1 Pendekatan dan Teknik Audit:', bold=True)
add_para('• Wawancara: Kepada Super Admin, Staf Kasir, dan Teknisi LKTech.\n• Observasi Sistem: Uji penetrasi ringan (black-box testing) pada antarmuka login, pengujian idle session, dan peninjauan manajemen Activity Log.\n• Review Kode: Peninjauan konfigurasi config/session.php, implementasi middleware perlindungan rute, dan prosedur database dump.', alignment='justify')

add_para('4.2 Model Penilaian Capability (0–5):', bold=True)
add_para('Penilaian dilakukan menggunakan validasi dokumen dan kondisi operasional. Bukti dinilai secara objektif untuk menetapkan level 0 hingga 5 sesuai pedoman COBIT 2019 Performance Management.', alignment='justify')

add_para('4.3 Skema Rating:', bold=True)
add_para('Fully (>85%), Largely (50–85%), Partially (15–50%), Not (<15%).', alignment='justify')

# 5. Gambaran Objek Audit (Sisfo)
add_heading('5. Gambaran Objek Audit (Sisfo)', level=2)
add_para('5.1 Deskripsi Singkat Sistem:', bold=True)
add_para('Sistem LKTech merupakan perangkat lunak ERP UMKM yang mensinergikan empat modul utama—inventori, penjualan (POS), penyewaan, dan servis—ke dalam platform web live. Sistem diproteksi oleh keamanan berlapis seperti Hash Bcrypt, pencegahan SQL Injection dengan PDO, dan proteksi CSRF.', alignment='justify')

add_para('5.2 Alur Proses Bisnis Utama:', bold=True)
add_para('Login Terproteksi 2FA → Dashboard Sesuai Role (RBAC) → Transaksi POS (Auto-Deduct Stok) / Pembaruan Status Servis → Perekaman di Activity Log.', alignment='justify')

# 6. Hasil Penilaian Capability
add_heading('6. Hasil Penilaian Capability (Per Proses/Praktik/Aktivitas)', level=2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['Proses/Praktik', 'Current', 'Target', 'Gap', 'Dokumen/Bukti Pendukung']
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], '1a3c5e')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].bold = True

rows_data = [
    ['DSS04 – Managed Continuity', '2', '3', '-1', 'Proses backup ada, tetapi hanya dilakukan secara insidental manual oleh admin.'],
    ['DSS05 – Managed Security Services', '3', '4', '-1', 'RBAC dan 2FA berjalan konsisten (Level 3 tercapai). Namun sesi tidak dihentikan otomatis saat idle.'],
    ['DSS05.02 – Manage endpoint security', '2', '3', '-1', 'Terminal kasir di area publik berisiko karena sistem tidak memiliki fungsi auto-logout.'],
    ['DSS05.04 – Manage user identity and logical access', '4', '4', '0', 'Manajemen akses sangat matang (RBAC & Wajib 2FA beroperasi penuh dan dicatat oleh Activity Log).']
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para()
add_para('Catatan Penilaian: Proses DSS04 terhenti di Level 2 (Managed) karena kebergantungan penuh pada intervensi manusia. Untuk mencapai Level 3 (Established), harus ada otomatisasi terjadwal. Sebaliknya, DSS05 menunjukkan kematangan yang sangat baik di area akses logika (Level 4 tercapai pada DSS05.04) berkat integrasi pustaka eksternal (2FA & RBAC Laravel). Namun praktik keamanan endpoint (terminal fisik PC Kasir) masih lemah (Level 2).', italic=True, alignment='justify')

# 7. Hasil Perhitungan Maturity
add_heading('7. Hasil Perhitungan Maturity', level=2)

table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
headers = ['Domain/Area', 'Metode', 'Nilai', 'Level', 'Catatan']
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], '1a3c5e')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].bold = True

rows_data2 = [
    ['DSS04 (Continuity)', 'Rata-rata Praktik', '2,00', 'Level 2', 'Butuh implementasi CRON Job'],
    ['DSS05 (Security Services)', 'Berbobot', '3,25', 'Level 3', 'Praktik autentikasi sangat kuat, manajemen sesi lemah'],
    ['Keseluruhan (Scope Audit)', 'Rata-rata', '2,62', 'Level 2', 'Kapabilitas operasional tinggi tetapi kelangsungan layanan perlu ditingkatkan']
]
for row_data in rows_data2:
    row_cells = table2.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para()
add_para('Analisis Maturity: Secara umum, kapabilitas keamanan (DSS05) berada pada Level 3, yang merupakan prestasi signifikan untuk kelas UMKM ritel. Hal ini didukung penuh oleh arsitektur framework Laravel yang menopang sistem. Sayangnya, celah pada ketahanan data historis (DSS04) menyeret nilai keseluruhan menjadi 2,62 (Level 2). Agar infrastruktur TI LKTech matang secara seimbang, perbaikan tata kelola otomatisasi infrastruktur server menjadi kunci utama menuju Target Maturity Level 3 di semua lini.', italic=True, alignment='justify')

# 8. Temuan Audit & Analisis Risiko
add_heading('8. Temuan Audit & Analisis Risiko', level=2)

table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Table Grid'
hdr_cells = table3.rows[0].cells
headers = ['No', 'Temuan', 'Risiko', 'Severity', 'Proses COBIT']
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], '1a3c5e')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].bold = True

rows_data3 = [
    ['1', 'Tidak ada algoritma auto-logout untuk sesi yang idle.', 'PC Kasir yang ditinggalkan dapat disalahgunakan pihak tidak bertanggung jawab.', 'High', 'DSS05.02'],
    ['2', 'Pencadangan Database SQL dilakukan secara manual.', 'Hilangnya data transaksi permanen (dampak fatal) jika server Cloud mengalami kegagalan teknis.', 'High', 'DSS04.07'],
    ['3', 'Activity Log tersedia dalam format tabel dasar.', 'Sulit mengidentifikasi anomali keamanan karena absennya representasi visual (grafik/chart).', 'Medium', 'MEA01.03']
]
for row_data in rows_data3:
    row_cells = table3.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para()

# 9. Analisis Gap
add_heading('9. Analisis Gap (Current vs Target)', level=2)
table4 = doc.add_table(rows=1, cols=6)
table4.style = 'Table Grid'
hdr_cells = table4.rows[0].cells
headers = ['Proses/Aktivitas', 'Current', 'Target', 'Gap', 'Dampak Bisnis', 'Prioritas']
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], '1a3c5e')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].bold = True

rows_data4 = [
    ['DSS04.07 – Manajemen Backup Data', '2', '3', '1', 'Operasional lumpuh total jika database utama korup.', '5 (Sangat Tinggi)'],
    ['DSS05.02 – Keamanan Perangkat (Session)', '2', '3', '1', 'Celah pencurian uang fisik di laci kasir atau manipulasi inventori.', '5 (Sangat Tinggi)'],
    ['MEA01.03 – Pengumpulan/Analisis Log', '2', '3', '1', 'Keterlambatan deteksi insiden (seperti brute-force login).', '3 (Sedang)']
]
for row_data in rows_data4:
    row_cells = table4.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para()

# 10. Rekomendasi
add_heading('10. Rekomendasi & Rencana Tindak Lanjut', level=2)
table5 = doc.add_table(rows=1, cols=7)
table5.style = 'Table Grid'
hdr_cells = table5.rows[0].cells
headers = ['No', 'Rekomendasi Teknis', 'Output / Deliverable', 'PIC', 'Timeline', 'Indikator Sukses', 'Prioritas']
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], '1a3c5e')
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].runs[0].bold = True

rows_data5 = [
    ['1', 'Aktivasi CRON Job Backup ke layanan pihak ketiga (contoh: Google Drive / S3).', 'Script dump otomatis yang tereksekusi harian jam 00:00.', 'Teknisi IT', 'Q3 2026', 'Backup file terbentuk otomatis setiap hari.', 'High'],
    ['2', 'Implementasi Batas Sesi (Auto-Logout). Modifikasi konfigurasi Laravel session.lifetime.', 'Sesi mati otomatis setelah 10-15 menit idle.', 'Teknisi IT', 'Q3 2026', 'Pengguna dipaksa login 2FA ulang jika layar dibiarkan.', 'High'],
    ['3', 'Pengembangan Visualisasi Activity Log.', 'Dashboard Analytic dengan grafik tren.', 'Teknisi IT', 'Q4 2026', 'Super Admin dapat memantau event secara visual.', 'Medium'],
    ['4', 'Penyusunan SOP Keamanan Karyawan.', 'Dokumen fisik SOP Penanganan Kredensial.', 'Super Admin', 'Q1 2027', 'Seluruh karyawan memahami SLA keamanan.', 'Medium']
]
for row_data in rows_data5:
    row_cells = table5.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para()

# 11. Roadmap Peningkatan
add_heading('11. Roadmap Peningkatan', level=2)
add_para('Analisis SWOT:', bold=True)
add_para('• Strengths: RBAC dan 2FA sudah beroperasi penuh dan melindungi sistem dengan sangat ketat (kematangan Level 4 di aspek logical access). Fitur pendataan (Activity Log) beroperasi transparan.\n• Weaknesses: Ketiadaan manajemen sesi perangkat fisik (kasir) yang mati secara otonom. Tidak berjalannya script cron job penjadwalan.\n• Opportunities: Penggunaan kerangka kerja modern (Laravel) membuat perbaikan pada session timeout dan pembuatan fitur visualisasi dashboard sangat mudah dilakukan tanpa mengubah arsitektur inti.\n• Threats: Serangan siber eksternal (Ransomware) pada server cloud berpotensi menghapus seluruh data jika tidak memiliki proteksi basis data harian (backup asinkron).', alignment='justify')
add_para('Roadmap Implementasi:', bold=True)
add_para('• Fase 1 (Jangka Pendek / Segera): Stabilisasi Keamanan Dasar. Aktivasi CRON Job Backup dan Penyesuaian config/session.php menjadi 10 Menit.\n• Fase 2 (Jangka Menengah / 3-6 bulan): Pengembangan Visibilitas Audit. Pembuatan halaman Dashboard Monitoring grafik keamanan khusus peran Super Admin.\n• Fase 3 (Jangka Panjang / 6-12 bulan): Kematangan Tata Kelola. Rutinisasi audit keamanan berkala dan kewajiban pelatihan SOP bagi pengguna akhir sistem LKTech.', alignment='justify')

# 12. Kesimpulan
add_heading('12. Kesimpulan', level=2)
add_para('• Ringkasan Capability & Maturity: Sistem Informasi LKTech mencatatkan performa keamanan akses logis (DSS05) di Level 3 (Established) yang luar biasa, namun operasional pencadangan kelangsungan bisnis (DSS04) tertahan di Level 2 (Managed).\n• Risiko Utama: Risiko terbesar yang menghantui operasional sistem saat ini adalah kegagalan fisik/perangkat lunak server (disaster) yang dapat menghapus data permanen karena ketiadaan proses pencadangan otonom, disusul oleh celah keamanan fisik akibat tidak adanya auto-logout.\n• Rekomendasi Prioritas: Otomatisasikan perlindungan data dengan mengeksekusi penulisan script CRON Job backup sekarang juga. Kedua, tutup celah keamanan terminal kasir dengan mengaktifkan penghentian sesi idle otomatis (maksimal 10-15 menit).', alignment='justify')

# 13. Lampiran
add_heading('13. Lampiran', level=2)
add_para('• Lampiran A. Daftar Responden (Super Admin, Kasir, Teknisi)\n• Lampiran B. Instrumen Kuesioner COBIT 2019 (Domain DSS04 & DSS05)\n• Lampiran C. Screenshot Middleware 2FA & Role Laravel\n• Lampiran D. Rekap Perhitungan Capability & Maturity LKTech')

doc.save(DST)
print("Docx generated successfully!")
