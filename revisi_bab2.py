"""
Script Revisi BAB 2 - Mursanto (19252417)
Revisi: Sub-bab 2.4 Penelitian Terdahulu
- Ubah format dari naratif per penelitian ke format tabel
- Tambahkan analisis penelitian terdahulu (hubungan, persamaan, perbedaan, kebaruan)
"""

import shutil
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# PATHS
# ============================================================
SRC = r'D:\Project\Kebutuhan Skripsi\SKRIPSI\BAB 2_MURSANTO_19252417_silakan direvisi lagi.docx'
DST = r'D:\Project\Kebutuhan Skripsi\SKRIPSI\BAB 2_REVISI_FINAL_MURSANTO.docx'

shutil.copy2(SRC, DST)
print(f"File disalin ke: {DST}")

doc = Document(DST)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_paragraph_format(para, left_indent=None, space_before=None, space_after=None, line_spacing=None):
    """Set paragraph formatting"""
    pPr = para._p.get_or_add_pPr()
    if left_indent is not None:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left_indent))
        pPr.append(ind)

def add_heading_style(para, level=1):
    """Apply heading style to paragraph"""
    if level == 1:
        para.style = doc.styles['Heading 1']
    elif level == 2:
        para.style = doc.styles['Heading 2']
    elif level == 3:
        para.style = doc.styles['Heading 3']

def set_cell_background(cell, hex_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_bold(cell, bold=True):
    """Make all text in cell bold"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold

def set_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2F4F8F')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_run_with_style(para, text, bold=False, italic=False, font_size=None, color=None):
    """Add a run with specific styling"""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# ============================================================
# FIND PARAGRAPHS INDEX for "Penelitian Terdahulu"
# ============================================================

print("\nMencari posisi 'Penelitian Terdahulu' di dokumen...")
penelitian_idx = None
end_of_bab2_idx = len(doc.paragraphs) - 1

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Penelitian Terdahulu' in text and ('2.4' in text or text == 'Penelitian Terdahulu'):
        penelitian_idx = i
        print(f"  Ditemukan di paragraf [{i}]: '{text}'")
        break
    elif 'Penelitian Terdahulu' in text:
        penelitian_idx = i
        print(f"  Ditemukan di paragraf [{i}]: '{text}'")
        break

if penelitian_idx is None:
    print("WARN: Mencari 'Penelitian' saja...")
    for i, para in enumerate(doc.paragraphs):
        if 'Penelitian' in para.text and 'Terdahulu' in para.text:
            penelitian_idx = i
            print(f"  Ditemukan di paragraf [{i}]: '{para.text[:80]}'")
            break

print(f"\nTotal paragraf: {len(doc.paragraphs)}")
print(f"Index Penelitian Terdahulu: {penelitian_idx}")

# Show context around penelitian_idx
if penelitian_idx:
    for i in range(max(0, penelitian_idx-2), min(len(doc.paragraphs), penelitian_idx+20)):
        print(f"  [{i}] {doc.paragraphs[i].text[:100]}")

# ============================================================
# REMOVE OLD CONTENT (from penelitian_idx+1 to end of BAB 2)
# ============================================================
# The paragraphs after "Penelitian Terdahulu" heading until end of document
# need to be replaced

if penelitian_idx is not None:
    # Get the XML body
    body = doc.element.body
    
    # Get all paragraph elements
    all_paras = body.findall('.//' + qn('w:p'))
    all_tables = body.findall('.//' + qn('w:tbl'))
    
    print(f"\nTotal elements in body: {len(list(body))}")
    
    # Find the paragraph element for penelitian_idx
    doc_paragraphs = doc.paragraphs
    target_para_elem = doc_paragraphs[penelitian_idx]._p
    
    # Remove all elements AFTER the penelitian_terdahulu heading
    body_children = list(body)
    found_target = False
    elements_to_remove = []
    
    for child in body_children:
        if found_target:
            # Remove this element (it's content after "Penelitian Terdahulu")
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if tag in ['p', 'tbl']:
                elements_to_remove.append(child)
        if child == target_para_elem:
            found_target = True
    
    print(f"Elements to remove after heading: {len(elements_to_remove)}")
    for elem in elements_to_remove:
        body.remove(elem)
    
    print("Konten lama 'Penelitian Terdahulu' berhasil dihapus.")

# ============================================================
# ADD NEW CONTENT: Paragraf Pengantar
# ============================================================

print("\nMenambahkan konten baru Penelitian Terdahulu...")

# Paragraf pengantar
intro_para = doc.add_paragraph()
intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = intro_para.add_run(
    'Kajian terhadap penelitian-penelitian terdahulu yang relevan menjadi pijakan penting dalam memposisikan penelitian ini di tengah lanskap keilmuan yang ada. '
    'Lima penelitian terdahulu berikut dipilih karena memiliki keterkaitan langsung dengan tema pengembangan sistem informasi berbasis Laravel, manajemen inventori dan penjualan, '
    'serta implementasi keamanan autentikasi (Tabel 2.1).'
)
run.font.size = Pt(12)

# Caption tabel
caption_para = doc.add_paragraph()
caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption_run = caption_para.add_run('Tabel 2.1 Ringkasan Penelitian Terdahulu')
caption_run.bold = True
caption_run.font.size = Pt(11)

# ============================================================
# ADD TABLE: Penelitian Terdahulu
# ============================================================

# Data penelitian terdahulu
penelitian_data = [
    {
        'no': '1',
        'peneliti': 'Widodo & Fithri\n(2025)',
        'judul': 'Implementasi Sistem Manajemen Inventory untuk Meningkatkan Efisiensi Stok Barang di PT Internusa Master Niaga',
        'sumber': 'PROFICIO: Jurnal Pengabdian Kepada Masyarakat, 6(1), 720',
        'metode': 'Implementatif, Laravel Framework',
        'temuan': 'Sistem berbasis Laravel mampu menekan selisih fisik gudang dan mengotomatisasi pencatatan pergerakan barang yang sebelumnya manual.',
        'relevansi': 'Landasan teknologi dan justifikasi pemilihan Laravel sebagai framework utama dalam modul inventori LKtech.',
    },
    {
        'no': '2',
        'peneliti': 'Rahmatuloh, Kamil & Nirwan\n(2024)',
        'judul': 'Rancang Bangun Aplikasi Point of Sale Toko Fashion Berbasis Website',
        'sumber': 'Jurnal Teknik Informatika, 16(3), 127',
        'metode': 'Waterfall, berbasis web',
        'temuan': 'Sistem POS mampu mengurangi human error, mempercepat transaksi kasir, dan menghasilkan laporan otomatis.',
        'relevansi': 'Acuan pengembangan modul kasir (POS) dan pelaporan transaksi pada LKtech, termasuk auto-deduct stok.',
    },
    {
        'no': '3',
        'peneliti': 'Sukma & Solichin\n(2025)',
        'judul': 'Penerapan Autentikasi Dua Faktor Menggunakan TOTP Berbasis Email dan Google Authenticator',
        'sumber': 'Prosiding SENAFTI, 4(2), 164',
        'metode': 'Pengembangan dan implementasi 2FA-TOTP',
        'temuan': '2FA dengan TOTP terbukti meningkatkan keamanan dari ancaman brute-force tanpa mengorbankan usability.',
        'relevansi': 'Rujukan teknis utama integrasi library TOTP dengan Laravel dan mekanisme verifikasi QR Code pada LKtech.',
    },
    {
        'no': '4',
        'peneliti': 'Wau\n(2022)',
        'judul': 'Pengembangan Sistem Informasi Persediaan Gudang Berbasis Website dengan Metode Waterfall',
        'sumber': 'MAROSTEK: Jurnal Teknik, Komputer, Agroteknologi dan Sains, 1(1), 10–23',
        'metode': 'Waterfall (analisis, desain, implementasi, pengujian)',
        'temuan': 'Sistem Waterfall menghasilkan produk sesuai spesifikasi dan menggantikan proses manual yang rentan inkonsistensi data.',
        'relevansi': 'Landasan pemilihan metode Waterfall dan referensi konseptual komponen sistem informasi inventori.',
    },
    {
        'no': '5',
        'peneliti': 'Alvansyah, Kiswanto, Nasution & Zulfi\n(2025)',
        'judul': 'Implementasi Zero Trust Architecture dengan Multi-Factor Authentication dan Continuous Verification pada Sistem Login Berbasis Web',
        'sumber': 'Jurnal Ilmiah Komputasi, 4(2), 617–623',
        'metode': 'Zero Trust Architecture + MFA pada web',
        'temuan': 'Zero Trust + MFA efektif mendeteksi login anomali dan menurunkan insiden akses tidak sah secara signifikan.',
        'relevansi': 'Memperkuat justifikasi implementasi 2FA pada LKtech dan perspektif arsitektural keamanan berlapis.',
    },
]

# Create table: 7 columns
headers = ['No', 'Peneliti / Tahun', 'Judul Penelitian', 'Sumber', 'Metode', 'Temuan Utama', 'Relevansi dengan Penelitian Ini']
col_widths = [Cm(0.8), Cm(2.8), Cm(4.2), Cm(3.0), Cm(2.5), Cm(4.0), Cm(4.0)]

table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Normal Table'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for i, width in enumerate(col_widths):
    table.columns[i].width = width

# Header row
hdr_cells = table.rows[0].cells
for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    set_cell_background(cell, '2F4F8F')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
for row_data in penelitian_data:
    row_cells = table.add_row().cells
    values = [
        row_data['no'],
        row_data['peneliti'],
        row_data['judul'],
        row_data['sumber'],
        row_data['metode'],
        row_data['temuan'],
        row_data['relevansi'],
    ]
    for i, (cell, val) in enumerate(zip(row_cells, values)):
        cell.text = val
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.size = Pt(9)
        # Alternate row shading
        if int(row_data['no']) % 2 == 0:
            set_cell_background(cell, 'EEF2FF')

set_table_borders(table)
print("Tabel Penelitian Terdahulu berhasil dibuat.")

# ============================================================
# ADD SOURCE
# ============================================================
src_para = doc.add_paragraph('(Sumber: Olahan Peneliti, 2026)')
src_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in src_para.runs:
    run.italic = True
    run.font.size = Pt(10)

# ============================================================
# ANALISIS PENELITIAN TERDAHULU
# ============================================================

doc.add_paragraph()  # spacer

# Sub-heading: Analisis
analisis_heading = doc.add_paragraph()
analisis_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_h = analisis_heading.add_run('Analisis Penelitian Terdahulu')
run_h.bold = True
run_h.font.size = Pt(12)

# Paragraf analisis
analisis_text = doc.add_paragraph()
analisis_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
analisis_text.add_run(
    'Berdasarkan kajian terhadap kelima penelitian terdahulu yang dirangkum dalam Tabel 2.1 di atas, '
    'dapat diidentifikasi hubungan, persamaan, perbedaan, dan kebaruan yang membedakan penelitian ini dari penelitian sebelumnya.'
).font.size = Pt(12)

# Sub: Persamaan
persamaan_heading = doc.add_paragraph()
persamaan_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = persamaan_heading.add_run('a. Persamaan dengan Penelitian Sebelumnya')
r.bold = True
r.font.size = Pt(12)

persamaan_body = doc.add_paragraph()
persamaan_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
persamaan_body.add_run(
    'Penelitian ini memiliki beberapa persamaan mendasar dengan penelitian-penelitian terdahulu. '
    'Pertama, penelitian ini sama-sama menggunakan framework Laravel sebagai teknologi pengembangan utama, '
    'selaras dengan pendekatan yang digunakan oleh Widodo & Fithri (2025) dalam membangun sistem manajemen inventori. '
    'Kedua, penelitian ini mengadopsi metode Waterfall sebagai pendekatan pengembangan perangkat lunak, '
    'konsisten dengan Rahmatuloh et al. (2024) dan Wau (2022) yang terbukti menghasilkan sistem yang terstruktur dan sesuai spesifikasi. '
    'Ketiga, fokus pada pengurangan human error dan otomatisasi proses bisnis menjadi tujuan bersama yang ditemukan pula dalam penelitian Rahmatuloh et al. (2024). '
    'Keempat, implementasi mekanisme Two-Factor Authentication (2FA) berbasis TOTP dan Google Authenticator pada penelitian ini '
    'memiliki kesamaan teknis yang kuat dengan penelitian Sukma & Solichin (2025) dan Alvansyah et al. (2025).'
).font.size = Pt(12)

# Sub: Perbedaan
perbedaan_heading = doc.add_paragraph()
perbedaan_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = perbedaan_heading.add_run('b. Perbedaan dengan Penelitian Sebelumnya')
r2.bold = True
r2.font.size = Pt(12)

perbedaan_body = doc.add_paragraph()
perbedaan_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
perbedaan_body.add_run(
    'Meskipun terdapat kesamaan dalam beberapa aspek, penelitian ini memiliki sejumlah perbedaan signifikan dari penelitian-penelitian terdahulu. '
    'Penelitian Widodo & Fithri (2025) dan Rahmatuloh et al. (2024) berfokus pada satu aspek tunggal (inventori atau POS) tanpa mengintegrasikan modul keamanan autentikasi yang mendalam. '
    'Penelitian Wau (2022) hanya membahas sistem persediaan gudang sederhana tanpa mekanisme kontrol akses berbasis peran (RBAC) maupun 2FA. '
    'Sementara itu, penelitian Sukma & Solichin (2025) dan Alvansyah et al. (2025) berfokus pada aspek keamanan autentikasi secara parsial, '
    'tanpa mengintegrasikannya ke dalam sistem informasi bisnis yang lengkap mencakup inventori, penjualan, penyewaan, dan servis. '
    'Penelitian ini berbeda karena mengintegrasikan semua aspek tersebut dalam satu platform terpadu.'
).font.size = Pt(12)

# Sub: Kebaruan
kebaruan_heading = doc.add_paragraph()
kebaruan_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
r3 = kebaruan_heading.add_run('c. Kebaruan (Novelty) Penelitian')
r3.bold = True
r3.font.size = Pt(12)

kebaruan_body = doc.add_paragraph()
kebaruan_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kebaruan_body.add_run(
    'Kebaruan utama penelitian ini terletak pada integrasi holistik antara sistem informasi bisnis multi-modul (inventori, penjualan, penyewaan, servis) '
    'dengan mekanisme keamanan autentikasi berlapis yang komprehensif pada satu platform terpadu. '
    'Secara lebih spesifik, novelty penelitian ini mencakup: (1) integrasi wajib Two-Factor Authentication (2FA) berbasis Google Authenticator '
    'sebagai lapisan keamanan mandatory bagi seluruh pengguna internal, bukan opsional; '
    '(2) implementasi Role-Based Access Control (RBAC) dengan tiga hierarki peran (Super Admin, Kasir, Teknisi) yang diintegrasikan dengan middleware Laravel; '
    '(3) pencatatan Activity Log otomatis untuk seluruh aktivitas pengguna sebagai instrumen audit keamanan; '
    'dan (4) penerapan keamanan berlapis yang mencakup hash Bcrypt, CSRF token, proteksi XSS via Blade, '
    'dan pencegahan SQL Injection via Eloquent ORM, dalam sebuah sistem informasi UMKM yang bersifat end-to-end. '
    'Kombinasi ini belum ditemukan secara utuh dalam satu penelitian sebelumnya, sehingga penelitian ini memberikan kontribusi orisinal '
    'bagi pengembangan sistem informasi berbasis Laravel dengan penekanan pada keamanan autentikasi yang komprehensif.'
).font.size = Pt(12)

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f"\nSUKSES! BAB 2 REVISI berhasil disimpan ke:\n   {DST}")
print("\nPerubahan yang dilakukan:")
print("  - Sub-bab 2.4 Penelitian Terdahulu diubah ke format TABEL")
print("  - Ditambahkan: Analisis Persamaan dengan penelitian sebelumnya")
print("  - Ditambahkan: Analisis Perbedaan dengan penelitian sebelumnya")
print("  - Ditambahkan: Kebaruan (Novelty) penelitian ini")
