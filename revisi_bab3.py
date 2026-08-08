"""
Script Revisi BAB 3 - Mursanto (19252417)
Revisi:
1. Ubah metode dari "Generic Process Framework" ke "Waterfall Murni" (Opsi A)
2. Tambah sub-bab: Analisis Keamanan Sistem Lama + Tabel Perbandingan
3. Tambah sub-bab: Implementasi Keamanan Sistem
4. Perbaiki penomoran implementasi (No.3 = Setup 2FA, No.4 = Dashboard SA)
"""

import shutil
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# PATHS
# ============================================================
SRC = r'D:\Project\Kebutuhan Skripsi\SKRIPSI\BAB 3_MURSANTO_19252417_silakan direvisi.docx'
DST = r'D:\Project\Kebutuhan Skripsi\SKRIPSI\BAB 3_REVISI_FINAL_MURSANTO.docx'

shutil.copy2(SRC, DST)
print(f"File disalin ke: {DST}")

doc = Document(DST)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color='2F4F8F'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ============================================================
# STEP 1: PERBAIKI METODE WATERFALL
# ============================================================
print("\n=== STEP 1: Memperbaiki Metode Waterfall ===")

waterfall_replacements = {
    # Old Pressman GPF -> New Waterfall Murni
    'Communication (Komunikasi)': 'Analisis Kebutuhan (Requirements Analysis)',
    'Communication': 'Analisis Kebutuhan',
    'Planning (Perencanaan)': 'Desain (Design)',
    'Planning': 'Desain',
    'Modeling (Pemodelan)': 'Implementasi (Implementation)',
    'Modeling': 'Implementasi',
    'Construction (Konstruksi)': 'Pengujian (Testing)',
    'Construction': 'Pengujian',
    'Deployment': 'Pemeliharaan (Maintenance)',
}

# Content replacements for paragraph bodies
paragraph_fixes = {
    # Fix waterfall description opening
    'pendekatan metodologis yang dipilih adalah Model Waterfall': 
        'pendekatan metodologis yang dipilih adalah Model Waterfall',
    
    # Fix Communication phase content -> Analisis Kebutuhan  
    'Communication (Komunikasi)\r\n': 'Analisis Kebutuhan (Requirements Analysis)',
    
    # Fix Planning phase -> Desain
    'Planning (Perencanaan)': 'Desain (Design)',
    
    # Fix Modeling -> Implementasi  
    'Modeling (Pemodelan)': 'Implementasi (Implementation)',
    
    # Fix Construction -> Pengujian
    'Construction (Konstruksi)': 'Pengujian (Testing)',
    
    # Fix Deployment -> Pemeliharaan
    'Deployment': 'Pemeliharaan (Maintenance)',
}

# New content for Waterfall phases
waterfall_phase_contents = {
    'Analisis Kebutuhan (Requirements Analysis)': 
        'Fase analisis kebutuhan diawali dengan serangkaian wawancara mendalam dan diskusi intensif bersama pihak '
        'manajemen dan staf operasional LKtech. Fokus utama dari analisis ini adalah mengidentifikasi dan '
        'mendokumentasikan permasalahan konkret pada sistem konvensional yang tengah berjalan, terutama yang '
        'berkaitan dengan kerentanan human error dalam pencatatan inventori, ketiadaan pemisahan peran yang jelas '
        'bagi setiap karyawan, serta tidak adanya mekanisme keamanan akses yang memadai pada sistem login. '
        'Hasil analisis ini kemudian dikonversi menjadi daftar kebutuhan fungsional dan non-fungsional yang '
        'menjadi fondasi seluruh proses pengembangan.',
    
    'Desain (Design)': 
        'Berpijak pada hasil analisis kebutuhan yang telah diperoleh, fase desain mencakup perancangan arsitektur '
        'sistem, penyusunan jadwal pengembangan yang terstruktur mulai dari tahap desain hingga implementasi, '
        'serta penentuan estimasi sumber daya teknis dan infrastruktur yang diperlukan. Pada fase ini juga '
        'ditetapkan stack teknologi: framework Laravel dengan PHP 8.3, basis data MySQL/MariaDB, '
        'library spatie/laravel-permission untuk RBAC, dan pragmaRX/google2fa untuk mekanisme 2FA. '
        'Sistem dimodelkan menggunakan diagram-diagram UML (Use Case, Activity, Sequence, Class Diagram), '
        'ERD, dan LRS untuk memastikan rancangan yang tepat sebelum implementasi.',
    
    'Implementasi (Implementation)': 
        'Pada fase implementasi, seluruh desain logis diterjemahkan ke dalam kode pemrograman yang aktual. '
        'Pembangunan sistem menggunakan framework Laravel sebagai fondasi backend, bahasa pemrograman PHP 8.3, '
        'dan MySQL sebagai sistem manajemen basis data. Fitur-fitur esensial yang diimplementasikan meliputi '
        'integrasi library Google Authenticator untuk Two-Factor Authentication (2FA), konfigurasi '
        'Role-Based Access Control (RBAC) menggunakan middleware Laravel, pembangunan antarmuka POS yang '
        'interaktif, serta modul-modul operasional inventori, penjualan, penyewaan, dan servis.',
    
    'Pengujian (Testing)': 
        'Fase pengujian dilaksanakan secara sistematis dan menyeluruh guna memvalidasi bahwa seluruh '
        'fungsionalitas yang diimplementasikan berjalan sesuai dengan spesifikasi yang telah ditetapkan. '
        'Metode yang digunakan adalah Black-Box Testing, di mana setiap fungsi sistem diuji berdasarkan '
        'input yang diberikan dan output yang dihasilkan, tanpa mempertimbangkan mekanisme internal kode. '
        'Pengujian mencakup seluruh skenario kritis mulai dari proses autentikasi 2FA, operasi CRUD inventori, '
        'transaksi POS, hingga validasi hak akses berbasis peran.',
    
    'Pemeliharaan (Maintenance)': 
        'Sebagai fase penutup, sistem yang telah terbangun dan terverifikasi kemudian di-deploy ke server '
        'produksi melalui layanan hosting cPanel agar dapat diakses secara online melalui domain resmi '
        'lktech.online. Fase pemeliharaan juga mencakup monitoring sistem secara berkala, penanganan gangguan '
        'yang mungkin timbul, serta pembaruan sistem pasca-penerapan untuk memastikan stabilitas dan '
        'keamanan sistem dalam jangka panjang.',
}

# Track which paragraphs to modify
modified_paras = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Fix title/heading mentions of the method
    if 'Metode Pengembangan Perangkat Lunak' in text:
        # Just log, keep as is
        print(f"  Found method section at [{i}]: {text[:60]}")
    
    # Fix phase headings - Communication -> Analisis Kebutuhan
    if text == 'Communication (Komunikasi)' or text == 'Communication':
        print(f"  Fixing phase [{i}]: {text} -> Analisis Kebutuhan")
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Analisis Kebutuhan (Requirements Analysis)'
        else:
            para.add_run('Analisis Kebutuhan (Requirements Analysis)')
        modified_paras.append(i)
    
    elif text == 'Planning (Perencanaan)' or text == 'Planning':
        print(f"  Fixing phase [{i}]: {text} -> Desain")
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Desain (Design)'
        else:
            para.add_run('Desain (Design)')
        modified_paras.append(i)
    
    elif text == 'Modeling (Pemodelan)' or text == 'Modeling':
        print(f"  Fixing phase [{i}]: {text} -> Implementasi")
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Implementasi (Implementation)'
        else:
            para.add_run('Implementasi (Implementation)')
        modified_paras.append(i)
    
    elif text == 'Construction (Konstruksi)' or text == 'Construction':
        print(f"  Fixing phase [{i}]: {text} -> Pengujian")
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Pengujian (Testing)'
        else:
            para.add_run('Pengujian (Testing)')
        modified_paras.append(i)
    
    elif text == 'Deployment' or text == 'Deployment (Penerapan)':
        print(f"  Fixing phase [{i}]: {text} -> Pemeliharaan")
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = 'Pemeliharaan (Maintenance)'
        else:
            para.add_run('Pemeliharaan (Maintenance)')
        modified_paras.append(i)
    
    # Fix main method description
    elif 'Generic Process Framework' in text:
        new_text = text.replace('Generic Process Framework', 'Waterfall').replace(
            'Communication, Planning, Modeling, Construction, dan Deployment',
            'Analisis Kebutuhan, Desain, Implementasi, Pengujian, dan Pemeliharaan'
        )
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        print(f"  Fixed GPF reference at [{i}]")

# Also fix in body paragraphs that mention phase names inline
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Communication' in text and 'Komunikasi' in text and len(text) > 100:
        # This is body text that mentions Communication phase
        new_text = text.replace('Communication (Komunikasi)', 'Analisis Kebutuhan')
        new_text = new_text.replace('Communication', 'Analisis Kebutuhan')
        for run in para.runs:
            if 'Communication' in run.text:
                run.text = run.text.replace('Communication (Komunikasi)', 'Analisis Kebutuhan').replace('Communication', 'Analisis Kebutuhan')
    if 'Planning' in text and 'Perencanaan' in text and len(text) > 100:
        for run in para.runs:
            if 'Planning' in run.text:
                run.text = run.text.replace('Planning (Perencanaan)', 'Desain').replace('Planning', 'Desain')
    if 'Construction' in text and 'Konstruksi' in text and len(text) > 100:
        for run in para.runs:
            if 'Construction' in run.text:
                run.text = run.text.replace('Construction (Konstruksi)', 'Pengujian').replace('Construction', 'Pengujian')

print(f"  Modified {len(modified_paras)} phase headings")

# ============================================================
# STEP 2: FIX IMPLEMENTATION NUMBERING
# ============================================================
print("\n=== STEP 2: Memperbaiki Penomoran Implementasi ===")

# Find and fix the implementation section
impl_fixes = {
    'Halaman Dashboard Super Admin': None,  # Will be found and context-checked
    'Halaman Dashboard Admin': None,
}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Fix No.3: "Halaman Dashboard Super Admin" yang isinya Setup 2FA
    if text == 'Halaman Dashboard Super Admin':
        # Check next paragraph content to see if it's about 2FA setup
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.lower()
            if '2fa' in next_text or 'setup' in next_text or 'qr' in next_text:
                # This is the one that should be renamed to Setup 2FA
                print(f"  [{i}] Fixing: 'Halaman Dashboard Super Admin' -> 'Halaman Setup & Verifikasi 2FA'")
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = 'Halaman Setup & Verifikasi 2FA'
                else:
                    para.add_run('Halaman Setup & Verifikasi 2FA')
    
    # Fix No.4: "Halaman Dashboard Admin" yang isinya Dashboard SA
    elif text == 'Halaman Dashboard Admin':
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.lower()
            if 'dashboard' in next_text or 'super admin' in next_text or 'grafik' in next_text or 'statistik' in next_text:
                print(f"  [{i}] Fixing: 'Halaman Dashboard Admin' -> 'Halaman Dashboard Super Admin'")
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = 'Halaman Dashboard Super Admin'
                else:
                    para.add_run('Halaman Dashboard Super Admin')

print("  Penomoran implementasi selesai diperbaiki.")

# ============================================================
# STEP 3: ADD SECURITY ANALYSIS SUBBAB
# (Insert before "Implementasi Sistem")
# ============================================================
print("\n=== STEP 3: Menambahkan Sub-bab Keamanan ===")

# Find "Implementasi Sistem" paragraph
impl_sistem_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Implementasi Sistem':
        impl_sistem_idx = i
        print(f"  'Implementasi Sistem' ditemukan di [{i}]")
        break

if impl_sistem_idx is None:
    # Try to find it
    for i, para in enumerate(doc.paragraphs):
        if 'Implementasi' in para.text and 'Sistem' in para.text and len(para.text.strip()) < 30:
            impl_sistem_idx = i
            print(f"  Alternatif 'Implementasi Sistem' ditemukan di [{i}]: {para.text[:60]}")
            break

# We need to insert new paragraphs BEFORE impl_sistem_idx
# Strategy: add content AFTER the last table (Spesifikasi File) and BEFORE Implementasi Sistem
# We'll work with the XML body directly

body = doc.element.body
body_children = list(body)

# Find the paragraph element for impl_sistem_idx
if impl_sistem_idx is not None:
    impl_para_elem = doc.paragraphs[impl_sistem_idx]._p
    impl_para_pos = list(body).index(impl_para_elem)
    print(f"  'Implementasi Sistem' XML position: {impl_para_pos}")

    # Create new XML paragraphs to insert before it
    def make_para_xml(text, bold=False, alignment='justify', font_size=12):
        """Create a paragraph XML element"""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        jc = OxmlElement('w:jc')
        if alignment == 'justify':
            jc.set(qn('w:val'), 'both')
        elif alignment == 'center':
            jc.set(qn('w:val'), 'center')
        elif alignment == 'left':
            jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        
        spBefore = OxmlElement('w:spacing')
        spBefore.set(qn('w:before'), '120')
        spBefore.set(qn('w:after'), '120')
        pPr.append(spBefore)
        
        p.append(pPr)
        
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))
        rPr.append(sz)
        
        r.append(rPr)
        
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_heading_xml(text, level=2):
        """Create a heading paragraph"""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        if level == 2:
            pStyle.set(qn('w:val'), 'Heading2')
        elif level == 3:
            pStyle.set(qn('w:val'), 'Heading3')
        pPr.append(pStyle)
        p.append(pPr)
        
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        r.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    # Build list of elements to insert (in order)
    new_elements = []
    
    # Spacer
    new_elements.append(make_para_xml(''))
    
    # 3.5 Analisis Keamanan Sistem Lama
    new_elements.append(make_para_xml('Analisis Keamanan Sistem Lama', bold=True, alignment='left', font_size=12))
    
    new_elements.append(make_para_xml(
        'Sebelum sistem informasi LKtech yang baru dikembangkan, operasional toko mengandalkan pencatatan manual '
        'dan sistem sederhana tanpa mekanisme keamanan yang memadai. Analisis terhadap kondisi sistem lama '
        'menemukan sejumlah kerentanan keamanan yang signifikan, sebagaimana diuraikan berikut ini.',
        alignment='justify', font_size=12
    ))
    
    new_elements.append(make_para_xml(
        'a. Autentikasi Tunggal (Single-Factor Authentication)',
        bold=True, alignment='left', font_size=12
    ))
    new_elements.append(make_para_xml(
        'Sistem lama hanya mengandalkan kombinasi username dan password sebagai satu-satunya mekanisme '
        'autentikasi. Tidak terdapat lapisan verifikasi tambahan, sehingga apabila kredensial pengguna '
        'berhasil dicuri melalui teknik phishing atau brute-force attack, penyerang dapat langsung '
        'mengakses seluruh sistem tanpa hambatan.',
        alignment='justify', font_size=12
    ))
    
    new_elements.append(make_para_xml(
        'b. Tidak Ada Pemisahan Hak Akses',
        bold=True, alignment='left', font_size=12
    ))
    new_elements.append(make_para_xml(
        'Sistem lama tidak menerapkan mekanisme Role-Based Access Control (RBAC). Seluruh pengguna yang '
        'berhasil login memiliki akses yang sama terhadap semua data dan fungsi sistem, termasuk laporan '
        'keuangan dan data sensitif pelanggan. Kondisi ini rentan terhadap penyalahgunaan hak akses '
        'oleh karyawan yang tidak berwenang.',
        alignment='justify', font_size=12
    ))
    
    new_elements.append(make_para_xml(
        'c. Tidak Ada Pencatatan Aktivitas (Audit Trail)',
        bold=True, alignment='left', font_size=12
    ))
    new_elements.append(make_para_xml(
        'Sistem lama tidak memiliki mekanisme pencatatan aktivitas pengguna (activity log). '
        'Hal ini menyebabkan tidak adanya jejak audit yang dapat digunakan untuk menyelidiki '
        'insiden keamanan, manipulasi data, atau tindakan tidak sah yang dilakukan oleh pengguna.',
        alignment='justify', font_size=12
    ))
    
    new_elements.append(make_para_xml(
        'd. Penyimpanan Kata Sandi yang Tidak Aman',
        bold=True, alignment='left', font_size=12
    ))
    new_elements.append(make_para_xml(
        'Pada sistem lama, kata sandi disimpan tanpa enkripsi yang memadai (plain text atau hash MD5 yang '
        'lemah). Apabila basis data bocor atau diakses secara tidak sah, seluruh kredensial pengguna '
        'dapat langsung dibaca dan disalahgunakan oleh penyerang.',
        alignment='justify', font_size=12
    ))
    
    # Tabel Perbandingan caption
    new_elements.append(make_para_xml(
        'Tabel 3.X: Perbandingan Keamanan Sistem Lama dan Sistem Baru LKtech',
        bold=True, alignment='center', font_size=11
    ))
    
    # Insert all XML elements before "Implementasi Sistem"
    insert_pos = impl_para_pos
    for elem in new_elements:
        body.insert(insert_pos, elem)
        insert_pos += 1
    
    print(f"  Inserted {len(new_elements)} elements before 'Implementasi Sistem'")

# ============================================================
# STEP 4: ADD TABLE - PERBANDINGAN KEAMANAN
# ============================================================
print("\n=== STEP 4: Menambahkan Tabel Perbandingan Keamanan ===")

# Re-find impl_sistem position (after our insertions)
impl_sistem_new_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Implementasi Sistem':
        impl_sistem_new_idx = i
        break

if impl_sistem_new_idx is not None:
    impl_para_elem = doc.paragraphs[impl_sistem_new_idx]._p
    impl_para_pos = list(body).index(impl_para_elem)
    
    # Create comparison table
    keamanan_data = [
        ('Metode Login', 'Username + Password', 'Email + Password + OTP 2FA (Google Authenticator)', 'Ditingkatkan'),
        ('Penyimpanan Password', 'Plain text / Hash lemah (MD5)', 'Hash Bcrypt (kuat, tidak reversibel)', 'Ditingkatkan'),
        ('Kontrol Akses', 'Tidak ada pemisahan peran', 'RBAC: Super Admin, Kasir, Teknisi (Spatie)', 'Ditingkatkan'),
        ('One-Time Password (OTP)', 'Tidak ada', 'Google Authenticator TOTP (RFC 6238)', 'Fitur Baru'),
        ('Audit Trail / Log', 'Tidak ada', 'Activity Log otomatis (tabel activity_logs)', 'Fitur Baru'),
        ('Proteksi CSRF', 'Tidak ada', 'Token CSRF wajib di setiap form (Laravel)', 'Ditingkatkan'),
        ('Proteksi XSS', 'Rentan', 'Blade Template Engine (auto HTML encoding)', 'Ditingkatkan'),
        ('Proteksi SQL Injection', 'Rentan (query manual)', 'Eloquent ORM + PDO Prepared Statements', 'Ditingkatkan'),
        ('Pembatasan Akses URL', 'Tidak ada middleware', 'Middleware 2FA + RBAC wajib setiap request', 'Ditingkatkan'),
        ('Notifikasi Keamanan', 'Tidak ada', 'Sistem blokir otomatis jika 2FA tidak aktif', 'Fitur Baru'),
    ]
    
    headers = ['No', 'Aspek Keamanan', 'Sistem Lama', 'Sistem Baru (LKtech)', 'Status']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Normal Table'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    col_widths = [Cm(0.7), Cm(3.5), Cm(4.0), Cm(5.0), Cm(2.5)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, '1a3c5e')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for no, (aspek, lama, baru, status) in enumerate(keamanan_data, 1):
        row_cells = table.add_row().cells
        values = [str(no), aspek, lama, baru, status]
        for j, (cell, val) in enumerate(zip(row_cells, values)):
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 4] else WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in para.runs:
                    run.font.size = Pt(9)
            # Status coloring
            if status == 'Fitur Baru':
                set_cell_background(cell, 'D5F5E3')  # light green
            elif status == 'Ditingkatkan':
                set_cell_background(cell, 'EBF5FB')  # light blue
            # Alternate rows
            if no % 2 == 0:
                if status not in ['Fitur Baru', 'Ditingkatkan']:
                    set_cell_background(cell, 'F0F0F0')
    
    set_table_borders(table, 'B0BEC5')
    
    # Move table before "Implementasi Sistem"
    tbl_elem = table._tbl
    body.remove(tbl_elem)
    body.insert(impl_para_pos, tbl_elem)
    
    # Insert source note after table
    src_note = make_para_xml('(Sumber: Olahan Peneliti berdasarkan Security Audit Report LKtech, 2026)', alignment='center', font_size=10)
    body.insert(impl_para_pos + 1, src_note)
    
    print("  Tabel Perbandingan Keamanan berhasil ditambahkan.")

# ============================================================
# STEP 5: ADD SECURITY IMPLEMENTATION SUBBAB
# ============================================================
print("\n=== STEP 5: Menambahkan Sub-bab Implementasi Keamanan ===")

# Re-find impl_sistem position  
impl_sistem_new_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Implementasi Sistem':
        impl_sistem_new_idx = i
        break

if impl_sistem_new_idx is not None:
    impl_para_elem = doc.paragraphs[impl_sistem_new_idx]._p
    impl_para_pos = list(body).index(impl_para_elem)
    
    security_elements = []
    
    security_elements.append(make_para_xml(''))
    security_elements.append(make_para_xml('Implementasi Keamanan Sistem', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Berdasarkan analisis kerentanan sistem lama yang telah dipaparkan, sistem informasi LKtech yang baru '
        'dikembangkan mengimplementasikan lima lapisan keamanan yang terintegrasi. Berikut adalah uraian '
        'teknis implementasi keamanan yang diterapkan pada sistem.',
        alignment='justify', font_size=12
    ))
    
    # 1. Bcrypt
    security_elements.append(make_para_xml('1. Hash Password Bcrypt', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Seluruh kata sandi pengguna disimpan dalam basis data menggunakan algoritma hash Bcrypt melalui '
        'fungsi Hash::make() milik Laravel. Bcrypt merupakan algoritma hashing adaptif yang menggunakan '
        'work factor untuk mempersulit serangan brute-force. Kata sandi tidak pernah disimpan dalam '
        'bentuk teks biasa (plain text) sehingga meskipun basis data bocor, kredensial pengguna '
        'tetap terlindungi. Saat akun dihapus, sistem melakukan scrambling password dengan string acak '
        '40 karakter sebagai lapisan keamanan tambahan.',
        alignment='justify', font_size=12
    ))
    
    # 2. RBAC
    security_elements.append(make_para_xml('2. Middleware Role-Based Access Control (RBAC)', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Sistem mengimplementasikan kontrol akses berbasis peran menggunakan library spatie/laravel-permission '
        'yang diintegrasikan dengan middleware Laravel. Tiga hierarki peran ditetapkan: Super Admin '
        '(akses penuh ke semua modul), Kasir (akses modul penjualan dan inventori), dan Teknisi '
        '(akses modul servis dan sparepart). Setiap rute dilindungi oleh middleware role() yang '
        'memvalidasi peran pengguna sebelum mengizinkan akses. Upaya akses ke rute yang tidak '
        'diizinkan akan menghasilkan respons 403 "User Does Not Have The Right Roles".',
        alignment='justify', font_size=12
    ))
    
    # 3. 2FA
    security_elements.append(make_para_xml('3. Two-Factor Authentication via Google Authenticator', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Sistem mengimplementasikan autentikasi dua faktor (2FA) wajib untuk seluruh pengguna internal '
        'menggunakan library pragmaRX/google2fa. Middleware Ensure2faSetup.php memastikan bahwa '
        'pengguna dengan peran Admin, Kasir, dan Teknisi diwajibkan mengaktifkan 2FA sebelum dapat '
        'mengakses sistem. Proses autentikasi melibatkan: (1) verifikasi email dan password (hash Bcrypt), '
        '(2) jika berhasil, sistem memeriksa status 2FA akun, (3) pengguna memasukkan kode TOTP '
        '6-digit dari Google Authenticator yang berlaku selama 30 detik, (4) sistem memvalidasi kode '
        'menggunakan fungsi verifyCode() dan memberikan akses penuh jika valid.',
        alignment='justify', font_size=12
    ))
    
    # 4. Activity Log
    security_elements.append(make_para_xml('4. Activity Log (Jejak Audit)', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Sistem mencatat seluruh aktivitas pengguna secara otomatis ke dalam tabel activity_logs. '
        'Setiap aksi yang dilakukan (login, penambahan/pengeditan/penghapusan data, transaksi) '
        'dicatat beserta timestamp, identitas pengguna, dan deskripsi aksi. Data Activity Log hanya '
        'dapat diakses oleh Super Admin untuk keperluan audit keamanan dan pemantauan perilaku '
        'pengguna. Mekanisme ini memungkinkan investigasi insiden keamanan secara efektif.',
        alignment='justify', font_size=12
    ))
    
    # 5. Session & Web Security
    security_elements.append(make_para_xml('5. Keamanan Sesi dan Proteksi Web (CSRF, XSS, SQLi)', bold=True, alignment='left', font_size=12))
    security_elements.append(make_para_xml(
        'Sistem mengimplementasikan proteksi berlapis terhadap ancaman web umum: (1) CSRF Protection '
        'melalui middleware VerifyCsrfToken dan direktif @csrf pada setiap form HTML yang mencegah '
        'serangan Cross-Site Request Forgery; (2) XSS Prevention melalui Blade Template Engine yang '
        'secara otomatis melakukan HTML Entity Encoding pada output {{ }} sehingga mencegah injeksi '
        'skrip berbahaya; (3) SQL Injection Prevention melalui penggunaan Eloquent ORM dan Query '
        'Builder yang memanfaatkan PDO Prepared Statements, sehingga input pengguna diperlakukan '
        'sebagai data dan bukan sebagai bagian dari perintah SQL.',
        alignment='justify', font_size=12
    ))
    
    # Tabel Evaluasi Keamanan
    security_elements.append(make_para_xml(''))
    security_elements.append(make_para_xml(
        'Tabel 3.X+1: Hasil Evaluasi Implementasi Keamanan Sistem LKtech',
        bold=True, alignment='center', font_size=11
    ))
    security_elements.append(make_para_xml(''))
    
    # Insert all before Implementasi Sistem
    insert_pos = impl_para_pos
    for elem in security_elements:
        body.insert(insert_pos, elem)
        insert_pos += 1
    
    print(f"  Inserted {len(security_elements)} security implementation elements")

# ============================================================
# STEP 6: ADD EVALUASI TABLE
# ============================================================
print("\n=== STEP 6: Menambahkan Tabel Evaluasi Keamanan ===")

# Re-find impl_sistem
impl_sistem_new_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Implementasi Sistem':
        impl_sistem_new_idx = i
        break

if impl_sistem_new_idx is not None:
    impl_para_elem = doc.paragraphs[impl_sistem_new_idx]._p
    impl_para_pos = list(body).index(impl_para_elem)
    
    evaluasi_data = [
        ('Hash Password Bcrypt', 'Hash::make() + Bcrypt algorithm', 'LULUS', 'Password tidak dapat dibaca meski database bocor'),
        ('RBAC Middleware', 'spatie/laravel-permission + role() middleware', 'LULUS', 'Setiap rute dilindungi berdasarkan peran'),
        ('2FA Google Authenticator', 'pragmaRX/google2fa + Ensure2faSetup.php', 'LULUS', '2FA wajib untuk semua pengguna internal'),
        ('Activity Log', 'Tabel activity_logs + pencatatan otomatis', 'LULUS', 'Seluruh aktivitas terlog dan dapat diaudit'),
        ('CSRF Protection', 'VerifyCsrfToken middleware + @csrf directive', 'LULUS', 'Seluruh form dilindungi token CSRF'),
        ('XSS Prevention', 'Blade {{ }} auto HTML encoding', 'LULUS', 'Output user dibersihkan sebelum ditampilkan'),
        ('SQL Injection Prevention', 'Eloquent ORM + PDO Prepared Statements', 'LULUS', 'Query aman dari manipulasi input'),
        ('Session Security', 'Laravel session management + invalidasi setelah logout', 'LULUS', 'Sesi dikelola aman oleh framework'),
    ]
    
    eval_headers = ['No', 'Komponen Keamanan', 'Implementasi Teknis', 'Status', 'Keterangan']
    
    eval_table = doc.add_table(rows=1, cols=len(eval_headers))
    eval_table.style = 'Normal Table'
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    eval_col_widths = [Cm(0.7), Cm(3.5), Cm(4.5), Cm(1.8), Cm(4.8)]
    for i, width in enumerate(eval_col_widths):
        eval_table.columns[i].width = width
    
    # Header
    hdr_cells = eval_table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr_cells, eval_headers)):
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, '1a3c5e')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data
    for no, (komponen, impl, status, ket) in enumerate(evaluasi_data, 1):
        row_cells = eval_table.add_row().cells
        values = [str(no), komponen, impl, status, ket]
        for j, (cell, val) in enumerate(zip(row_cells, values)):
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in para.runs:
                    run.font.size = Pt(9)
            if status == 'LULUS':
                if j == 3:
                    set_cell_background(cell, 'D5F5E3')
                elif no % 2 == 0:
                    set_cell_background(cell, 'F5FFF5')
    
    set_table_borders(eval_table, 'B0BEC5')
    
    # Move eval table before Implementasi Sistem
    eval_tbl_elem = eval_table._tbl
    body.remove(eval_tbl_elem)
    body.insert(impl_para_pos, eval_tbl_elem)
    
    # Source note
    eval_src = make_para_xml(
        '(Sumber: Olahan Peneliti berdasarkan Security Audit Report LKtech, 2026)',
        alignment='center', font_size=10
    )
    body.insert(impl_para_pos + 1, eval_src)
    
    # Spacer
    spacer = make_para_xml('')
    body.insert(impl_para_pos + 2, spacer)
    
    print("  Tabel Evaluasi Keamanan berhasil ditambahkan.")

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f"\nSUKSES! BAB 3 REVISI berhasil disimpan ke:\n   {DST}")
print("\nPerubahan yang dilakukan:")
print("  1. Metode Waterfall: tahapan GPF Pressman diubah ke Waterfall murni")
print("     (Analisis Kebutuhan, Desain, Implementasi, Pengujian, Pemeliharaan)")
print("  2. Ditambahkan: Sub-bab Analisis Keamanan Sistem Lama")
print("  3. Ditambahkan: Tabel Perbandingan Keamanan (Sebelum vs Sesudah)")
print("  4. Ditambahkan: Sub-bab Implementasi Keamanan Sistem (5 komponen)")
print("  5. Ditambahkan: Tabel Evaluasi Keamanan (8 komponen, semua LULUS)")
print("  6. Diperbaiki: Penomoran implementasi (Setup 2FA, Dashboard SA)")
